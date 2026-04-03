#!/usr/bin/env python3
"""Create EPF Press Release on letterhead - PDF and Word versions"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.colors import HexColor
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from io import BytesIO
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
BASE_DIR = "/Users/macmini/clawd/repos/epf"
LETTERHEAD_PDF = os.path.join(BASE_DIR, "EPF_Letterhead_Blank.pdf")
LOGO_PATH = os.path.join(BASE_DIR, "assets/logos/epf-logo.png")
OUTPUT_PDF = os.path.join(BASE_DIR, "EPF_PressRelease_MothersDay_2026.pdf")
OUTPUT_DOCX = os.path.join(BASE_DIR, "EPF_PressRelease_MothersDay_2026.docx")

# Press release content
PRESS_RELEASE = """FOR IMMEDIATE RELEASE

Empowerment Path Foundation Launches Mother's Day Blessing Basket Initiative to Support Single Mothers Across Detroit and Metro Detroit

Detroit, MI — April 2, 2026

This Mother's Day, Empowerment Path Foundation is calling on the community to come together in support of single mothers through its Mother's Day Blessing Basket Initiative, a heartfelt effort designed to uplift, honor, and provide essential care items to mothers in need.

The initiative will provide 100 baskets for 100 moms — thoughtfully curated baskets filled with self-care items, essentials, and uplifting gifts for mothers participating in local programs across Detroit and Metro Detroit. Confirmed partner organizations include Peggy's Place, CASS Community Social Services/Fox Family Shelter, and Alternatives for Girls, all of which serve women and families navigating housing instability and life transitions.

Community members are invited to "Sponsor a Mom" for $100, which fully funds one blessing basket. Donations are being accepted now through April 25, 2026, with additional contributions welcomed in person during the basket assembly event. All donations are tax-deductible as allowed by law through our fiscal sponsor, Bringing Back the Block.

"We believe every mother deserves to feel seen, appreciated, and supported, especially those navigating life's challenges on their own," said Tina N. Bowden, Founder of Empowerment Path Foundation. "This initiative is about more than gifts. It is about restoring dignity, spreading love, and reminding these women that their community stands with them."

Event Details:
• Date: May 2, 2026
• Time: 2:00 p.m. – 8:00 p.m.
• Location: Lockeroom Lounge, 18290 Livernois Avenue, Detroit, Michigan 48221
• Purpose: Drop-offs, donations, and basket assembly

Volunteers, donors, and community supporters are encouraged to attend.

To ensure privacy and respect, participating mothers will not be publicly identified. When needed, identifiers such as nicknames or numbers will be used internally to personalize baskets while maintaining confidentiality.

Empowerment Path Foundation is actively seeking:
• Individual sponsors
• Corporate sponsors
• Donation partners
• Volunteers

This initiative is part of the organization's broader mission to support and empower single mothers and their children through resources, community programming, and advocacy.

To sponsor a mother, donate, or get involved, please contact:
Tina N. Bowden, Founder
Empowerment Path Foundation
empowermentpathfoundation@gmail.com
Phone: 313-937-6077
EmpowermentPathFoundation.org

About Empowerment Path Foundation
Empowerment Path Foundation is a community-driven nonprofit organization dedicated to supporting single mothers/fathers and their children through empowerment programs, resources, and strategic partnerships that promote stability, self-sufficiency, and long-term success.

Empowerment Path Foundation operates under fiscal sponsorship with Bringing Back the Block, a nonprofit organization supporting community-based initiatives. Donations made to Empowerment Path Foundation are tax-deductible to the extent allowed by law through Bringing Back the Block."""


def create_pdf():
    """Create PDF with letterhead"""
    # Create content PDF
    packet = BytesIO()
    
    # Use reportlab to create content
    c = canvas.Canvas(packet, pagesize=letter)
    width, height = letter
    
    # EPF brand colors
    navy = HexColor('#1a2a4a')
    
    # Start below letterhead header (assume ~1.5 inches for logo area)
    y_position = height - 1.8 * inch
    left_margin = 0.75 * inch
    right_margin = width - 0.75 * inch
    text_width = right_margin - left_margin
    
    # Set font
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(navy)
    
    # FOR IMMEDIATE RELEASE
    c.drawString(left_margin, y_position, "FOR IMMEDIATE RELEASE")
    y_position -= 0.35 * inch
    
    # Title
    c.setFont("Helvetica-Bold", 13)
    title = "Empowerment Path Foundation Launches Mother's Day Blessing Basket"
    c.drawString(left_margin, y_position, title)
    y_position -= 0.2 * inch
    title2 = "Initiative to Support Single Mothers Across Detroit and Metro Detroit"
    c.drawString(left_margin, y_position, title2)
    y_position -= 0.35 * inch
    
    # Date line
    c.setFont("Helvetica-Bold", 10)
    c.drawString(left_margin, y_position, "Detroit, MI — April 2, 2026")
    y_position -= 0.3 * inch
    
    # Body text
    c.setFont("Helvetica", 9.5)
    c.setFillColor(HexColor('#333333'))
    
    paragraphs = [
        "This Mother's Day, Empowerment Path Foundation is calling on the community to come together in support of single mothers through its Mother's Day Blessing Basket Initiative, a heartfelt effort designed to uplift, honor, and provide essential care items to mothers in need.",
        
        "The initiative will provide 100 baskets for 100 moms — thoughtfully curated baskets filled with self-care items, essentials, and uplifting gifts for mothers participating in local programs across Detroit and Metro Detroit. Confirmed partner organizations include Peggy's Place, CASS Community Social Services/Fox Family Shelter, and Alternatives for Girls, all of which serve women and families navigating housing instability and life transitions.",
        
        'Community members are invited to "Sponsor a Mom" for $100, which fully funds one blessing basket. Donations are being accepted now through April 25, 2026, with additional contributions welcomed in person during the basket assembly event. All donations are tax-deductible as allowed by law through our fiscal sponsor, Bringing Back the Block.',
        
        '"We believe every mother deserves to feel seen, appreciated, and supported, especially those navigating life\'s challenges on their own," said Tina N. Bowden, Founder of Empowerment Path Foundation. "This initiative is about more than gifts. It is about restoring dignity, spreading love, and reminding these women that their community stands with them."',
    ]
    
    from reportlab.lib.utils import simpleSplit
    
    for para in paragraphs:
        lines = simpleSplit(para, "Helvetica", 9.5, text_width)
        for line in lines:
            c.drawString(left_margin, y_position, line)
            y_position -= 0.16 * inch
        y_position -= 0.1 * inch
    
    # Event Details
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(navy)
    c.drawString(left_margin, y_position, "Event Details:")
    y_position -= 0.2 * inch
    
    c.setFont("Helvetica", 9.5)
    c.setFillColor(HexColor('#333333'))
    event_items = [
        "• Date: May 2, 2026",
        "• Time: 2:00 p.m. – 8:00 p.m.",
        "• Location: Lockeroom Lounge, 18290 Livernois Avenue, Detroit, Michigan 48221",
        "• Purpose: Drop-offs, donations, and basket assembly"
    ]
    for item in event_items:
        c.drawString(left_margin + 0.1*inch, y_position, item)
        y_position -= 0.16 * inch
    y_position -= 0.1 * inch
    
    c.drawString(left_margin, y_position, "Volunteers, donors, and community supporters are encouraged to attend.")
    y_position -= 0.25 * inch
    
    privacy_text = "To ensure privacy and respect, participating mothers will not be publicly identified. When needed, identifiers such as nicknames or numbers will be used internally to personalize baskets while maintaining confidentiality."
    lines = simpleSplit(privacy_text, "Helvetica", 9.5, text_width)
    for line in lines:
        c.drawString(left_margin, y_position, line)
        y_position -= 0.16 * inch
    y_position -= 0.1 * inch
    
    # Seeking section
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(navy)
    c.drawString(left_margin, y_position, "Empowerment Path Foundation is actively seeking:")
    y_position -= 0.2 * inch
    
    c.setFont("Helvetica", 9.5)
    c.setFillColor(HexColor('#333333'))
    for item in ["• Individual sponsors", "• Corporate sponsors", "• Donation partners", "• Volunteers"]:
        c.drawString(left_margin + 0.1*inch, y_position, item)
        y_position -= 0.16 * inch
    y_position -= 0.1 * inch
    
    mission_text = "This initiative is part of the organization's broader mission to support and empower single mothers and their children through resources, community programming, and advocacy."
    lines = simpleSplit(mission_text, "Helvetica", 9.5, text_width)
    for line in lines:
        c.drawString(left_margin, y_position, line)
        y_position -= 0.16 * inch
    y_position -= 0.15 * inch
    
    # Contact section
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(navy)
    c.drawString(left_margin, y_position, "To sponsor a mother, donate, or get involved, please contact:")
    y_position -= 0.2 * inch
    
    c.setFont("Helvetica", 9.5)
    c.setFillColor(HexColor('#333333'))
    contact_lines = [
        "Tina N. Bowden, Founder",
        "Empowerment Path Foundation",
        "empowermentpathfoundation@gmail.com",
        "Phone: 313-937-6077",
        "EmpowermentPathFoundation.org"
    ]
    for line in contact_lines:
        c.drawString(left_margin, y_position, line)
        y_position -= 0.15 * inch
    y_position -= 0.15 * inch
    
    # About section
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(navy)
    c.drawString(left_margin, y_position, "About Empowerment Path Foundation")
    y_position -= 0.2 * inch
    
    c.setFont("Helvetica", 9)
    c.setFillColor(HexColor('#333333'))
    about_text = "Empowerment Path Foundation is a community-driven nonprofit organization dedicated to supporting single mothers/fathers and their children through empowerment programs, resources, and strategic partnerships that promote stability, self-sufficiency, and long-term success."
    lines = simpleSplit(about_text, "Helvetica", 9, text_width)
    for line in lines:
        c.drawString(left_margin, y_position, line)
        y_position -= 0.14 * inch
    y_position -= 0.08 * inch
    
    fiscal_text = "Empowerment Path Foundation operates under fiscal sponsorship with Bringing Back the Block, a nonprofit organization supporting community-based initiatives. Donations made to Empowerment Path Foundation are tax-deductible to the extent allowed by law through Bringing Back the Block."
    lines = simpleSplit(fiscal_text, "Helvetica", 9, text_width)
    for line in lines:
        c.drawString(left_margin, y_position, line)
        y_position -= 0.14 * inch
    
    c.save()
    
    # Merge with letterhead
    packet.seek(0)
    content_pdf = PdfReader(packet)
    letterhead_pdf = PdfReader(LETTERHEAD_PDF)
    
    output = PdfWriter()
    
    # Get letterhead page and merge content
    letterhead_page = letterhead_pdf.pages[0]
    content_page = content_pdf.pages[0]
    letterhead_page.merge_page(content_page)
    output.add_page(letterhead_page)
    
    with open(OUTPUT_PDF, "wb") as f:
        output.write(f)
    
    print(f"Created PDF: {OUTPUT_PDF}")


def create_docx():
    """Create Word document with letterhead"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add logo
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.5))
    
    doc.add_paragraph()  # Spacer
    
    # FOR IMMEDIATE RELEASE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("FOR IMMEDIATE RELEASE")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    # Title
    p = doc.add_paragraph()
    run = p.add_run("Empowerment Path Foundation Launches Mother's Day Blessing Basket Initiative to Support Single Mothers Across Detroit and Metro Detroit")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    # Date
    p = doc.add_paragraph()
    run = p.add_run("Detroit, MI — April 2, 2026")
    run.bold = True
    run.font.size = Pt(11)
    
    # Body paragraphs
    body_paras = [
        "This Mother's Day, Empowerment Path Foundation is calling on the community to come together in support of single mothers through its Mother's Day Blessing Basket Initiative, a heartfelt effort designed to uplift, honor, and provide essential care items to mothers in need.",
        
        "The initiative will provide 100 baskets for 100 moms — thoughtfully curated baskets filled with self-care items, essentials, and uplifting gifts for mothers participating in local programs across Detroit and Metro Detroit. Confirmed partner organizations include Peggy's Place, CASS Community Social Services/Fox Family Shelter, and Alternatives for Girls, all of which serve women and families navigating housing instability and life transitions.",
        
        'Community members are invited to "Sponsor a Mom" for $100, which fully funds one blessing basket. Donations are being accepted now through April 25, 2026, with additional contributions welcomed in person during the basket assembly event. All donations are tax-deductible as allowed by law through our fiscal sponsor, Bringing Back the Block.',
        
        '"We believe every mother deserves to feel seen, appreciated, and supported, especially those navigating life\'s challenges on their own," said Tina N. Bowden, Founder of Empowerment Path Foundation. "This initiative is about more than gifts. It is about restoring dignity, spreading love, and reminding these women that their community stands with them."',
    ]
    
    for para in body_paras:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)
    
    # Event Details
    p = doc.add_paragraph()
    run = p.add_run("Event Details:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    event_items = [
        ("Date:", "May 2, 2026"),
        ("Time:", "2:00 p.m. – 8:00 p.m."),
        ("Location:", "Lockeroom Lounge, 18290 Livernois Avenue, Detroit, Michigan 48221"),
        ("Purpose:", "Drop-offs, donations, and basket assembly")
    ]
    
    for label, value in event_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{label} ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)
    
    p = doc.add_paragraph("Volunteers, donors, and community supporters are encouraged to attend.")
    for run in p.runs:
        run.font.size = Pt(11)
    
    p = doc.add_paragraph("To ensure privacy and respect, participating mothers will not be publicly identified. When needed, identifiers such as nicknames or numbers will be used internally to personalize baskets while maintaining confidentiality.")
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Seeking section
    p = doc.add_paragraph()
    run = p.add_run("Empowerment Path Foundation is actively seeking:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    for item in ["Individual sponsors", "Corporate sponsors", "Donation partners", "Volunteers"]:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
    
    p = doc.add_paragraph("This initiative is part of the organization's broader mission to support and empower single mothers and their children through resources, community programming, and advocacy.")
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Contact section
    p = doc.add_paragraph()
    run = p.add_run("To sponsor a mother, donate, or get involved, please contact:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    contact_lines = [
        "Tina N. Bowden, Founder",
        "Empowerment Path Foundation",
        "empowermentpathfoundation@gmail.com",
        "Phone: 313-937-6077",
        "EmpowermentPathFoundation.org"
    ]
    for line in contact_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacer
    
    # About section
    p = doc.add_paragraph()
    run = p.add_run("About Empowerment Path Foundation")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    
    p = doc.add_paragraph("Empowerment Path Foundation is a community-driven nonprofit organization dedicated to supporting single mothers/fathers and their children through empowerment programs, resources, and strategic partnerships that promote stability, self-sufficiency, and long-term success.")
    for run in p.runs:
        run.font.size = Pt(10)
    
    p = doc.add_paragraph("Empowerment Path Foundation operates under fiscal sponsorship with Bringing Back the Block, a nonprofit organization supporting community-based initiatives. Donations made to Empowerment Path Foundation are tax-deductible to the extent allowed by law through Bringing Back the Block.")
    for run in p.runs:
        run.font.size = Pt(10)
    
    doc.save(OUTPUT_DOCX)
    print(f"Created Word doc: {OUTPUT_DOCX}")


if __name__ == "__main__":
    create_pdf()
    create_docx()
    print("\nDone! Files created:")
    print(f"  PDF:  {OUTPUT_PDF}")
    print(f"  Word: {OUTPUT_DOCX}")
